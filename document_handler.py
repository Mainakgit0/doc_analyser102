import os
import re
import hashlib
import logging
from typing import Tuple, Optional
import httpx
import asyncio
from datetime import datetime
from models import DocumentCache, PerformanceStats
from app import db

logger = logging.getLogger(__name__)


class DocumentHandler:
    """Enhanced document processing with caching and improved text extraction"""
    
    def __init__(self):
        self.timeout = httpx.Timeout(30.0)
    
    def generate_document_hash(self, content: bytes, identifier: str = "") -> str:
        """Generate unique hash for document content"""
        hasher = hashlib.sha256()
        hasher.update(content)
        if identifier:
            hasher.update(identifier.encode('utf-8'))
        return hasher.hexdigest()
    
    async def process_uploaded_file(self, file_content: bytes, filename: str) -> Tuple[str, str]:
        """Process uploaded file with enhanced text extraction"""
        try:
            # Generate document hash
            doc_hash = self.generate_document_hash(file_content, filename)
            
            # Check cache first
            cached_doc = DocumentCache.query.filter_by(document_hash=doc_hash).first()
            if cached_doc:
                cached_doc.last_accessed = datetime.utcnow()
                db.session.commit()
                
                # Update performance stats
                stats = PerformanceStats.get_stats()
                stats.update_stats(cache_hit=True)
                
                logger.info(f"Retrieved cached document: {filename}")
                return cached_doc.processed_text, doc_hash
            
            # Process the file
            text = await self._extract_text_from_bytes(file_content, filename)
            
            if not text.strip():
                raise ValueError("No text could be extracted from the uploaded file")
            
            # Cache the processed document
            self._cache_document(doc_hash, text, filename=filename)
            
            # Update performance stats
            stats = PerformanceStats.get_stats()
            is_pdf = filename.lower().endswith('.pdf')
            stats.update_stats(
                pdf_processed=is_pdf,
                chars_extracted=len(text)
            )
            
            logger.info(f"Processed uploaded file: {filename}, extracted {len(text)} characters")
            return text, doc_hash
            
        except Exception as e:
            logger.error(f"Error processing uploaded file {filename}: {e}")
            raise ValueError(f"Failed to process uploaded file: {str(e)}")
    
    async def download_and_process_document(self, url: str) -> Tuple[str, str]:
        """Download and process document from URL with enhanced extraction"""
        try:
            # Download the document
            async with httpx.AsyncClient(timeout=self.timeout) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                content = response.content
                content_type = response.headers.get('content-type', '').lower()
                
                # Generate document hash
                doc_hash = self.generate_document_hash(content, url)
                
                # Check cache first
                cached_doc = DocumentCache.query.filter_by(document_hash=doc_hash).first()
                if cached_doc:
                    cached_doc.last_accessed = datetime.utcnow()
                    db.session.commit()
                    
                    # Update performance stats
                    stats = PerformanceStats.get_stats()
                    stats.update_stats(cache_hit=True)
                    
                    logger.info(f"Retrieved cached document from URL: {url}")
                    return cached_doc.processed_text, doc_hash
                
                # Extract filename from URL
                filename = url.split('/')[-1] or 'document'
                
                # Process the content
                text = await self._extract_text_from_bytes(content, filename, content_type)
                
                if not text.strip():
                    raise ValueError("No text could be extracted from the document")
                
                # Cache the processed document
                self._cache_document(doc_hash, text, url)
                
                # Update performance stats
                stats = PerformanceStats.get_stats()
                is_pdf = 'pdf' in content_type or url.lower().endswith('.pdf')
                stats.update_stats(
                    pdf_processed=is_pdf,
                    chars_extracted=len(text)
                )
                
                logger.info(f"Downloaded and processed document from {url}, extracted {len(text)} characters")
                return text, doc_hash
                
        except Exception as e:
            logger.error(f"Error downloading document from {url}: {e}")
            raise ValueError(f"Failed to download and process document: {str(e)}")
    
    async def _extract_text_from_bytes(self, content: bytes, filename: str, 
                                     content_type: str = "") -> str:
        """Enhanced text extraction from various file formats"""
        filename_lower = filename.lower()
        
        # Determine file type
        is_pdf = ('pdf' in content_type or filename_lower.endswith('.pdf') or 
                 content.startswith(b'%PDF'))
        is_docx = ('docx' in content_type or filename_lower.endswith('.docx') or 
                  b'PK' in content[:4])
        is_doc = filename_lower.endswith('.doc')
        
        if is_pdf:
            return await self._extract_pdf_text(content)
        elif is_docx:
            return await self._extract_docx_text(content)
        elif is_doc:
            return await self._extract_doc_text(content)
        else:
            # Try as plain text
            return await self._extract_plain_text(content)
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Enhanced PDF text extraction"""
        try:
            import PyPDF2
            import io
            
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            
            logger.info(f"Processing PDF with {len(pdf_reader.pages)} pages")
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        # Enhanced text cleaning for PDFs
                        page_text = self._clean_pdf_text(page_text)
                        if page_text.strip():
                            text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                            logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    continue
            
            full_text = '\n\n'.join(text_parts)
            
            if not full_text.strip():
                logger.warning("No text extracted from PDF, attempting fallback")
                return content.decode('utf-8', errors='ignore')
            
            return full_text
            
        except ImportError:
            logger.warning("PyPDF2 not available, treating PDF as text")
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            # Fallback to raw content
            return content.decode('utf-8', errors='ignore')
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Enhanced DOCX text extraction"""
        try:
            from docx import Document
            import io
            
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
            
        except ImportError:
            logger.warning("python-docx not available, treating DOCX as text")
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            return content.decode('utf-8', errors='ignore')
    
    async def _extract_doc_text(self, content: bytes) -> str:
        """Extract text from old DOC format"""
        try:
            # For .doc files, we'd need additional libraries like python-docx2txt
            # For now, fall back to plain text extraction
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error processing DOC: {e}")
            return ""
    
    async def _extract_plain_text(self, content: bytes) -> str:
        """Extract plain text with encoding detection"""
        try:
            # Try common encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    # Basic validation - check if it's mostly readable text
                    if len([c for c in text[:1000] if c.isprintable() or c.isspace()]) > 800:
                        return text
                except UnicodeDecodeError:
                    continue
            
            # Final fallback
            return content.decode('utf-8', errors='ignore')
            
        except Exception as e:
            logger.error(f"Error extracting plain text: {e}")
            return ""
    
    def _clean_pdf_text(self, text: str) -> str:
        """Clean PDF-specific text artifacts"""
        if not text:
            return ""
        
        # Remove common PDF artifacts
        text = text.replace('\x00', '')  # Null characters
        text = text.replace('\uf0b7', '•')  # Bullet point characters
        text = text.replace('\uf020', ' ')  # Space characters
        
        # Fix hyphenated words split across lines
        text = re.sub(r'-\s*\n\s*', '', text)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix sentence spacing after periods
        text = re.sub(r'\.(?=[A-Z])', '. ', text)
        
        return text.strip()
    
    def _cache_document(self, doc_hash: str, text: str, url: str = None, 
                       filename: str = None):
        """Cache processed document"""
        try:
            # Clean text for database storage (remove NUL characters)
            clean_text = text.replace('\x00', '')
            
            cached_doc = DocumentCache()
            cached_doc.document_hash = doc_hash
            cached_doc.document_url = url
            cached_doc.document_filename = filename
            cached_doc.processed_text = clean_text
            cached_doc.text_length = len(clean_text)
            cached_doc.chunks_count = 0  # Will be updated when chunks are created
            
            db.session.add(cached_doc)
            db.session.commit()
            
            logger.info(f"Cached document with hash: {doc_hash}")
            
        except Exception as e:
            logger.error(f"Error caching document: {e}")
            db.session.rollback()


